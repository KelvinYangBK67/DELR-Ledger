#!/usr/bin/env python3
from __future__ import annotations

import csv
import calendar
import json
import re
from collections import OrderedDict
import sys
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
import tkinter as tk
from tkinter import filedialog, messagebox, ttk
from xml.etree.ElementTree import Element, SubElement, ElementTree, parse as xml_parse
from xml.sax.saxutils import escape

try:
    import yaml  # type: ignore
except Exception:
    yaml = None

try:
    from openpyxl import Workbook, load_workbook  # type: ignore
except Exception:
    Workbook = None
    load_workbook = None


try:
    from docx import Document  # type: ignore
except Exception:
    Document = None

try:
    from reportlab.lib import colors  # type: ignore
    from reportlab.lib.pagesizes import A4  # type: ignore
    from reportlab.lib.styles import getSampleStyleSheet  # type: ignore
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle  # type: ignore
    from reportlab.pdfbase import pdfmetrics  # type: ignore
    from reportlab.pdfbase.ttfonts import TTFont  # type: ignore
except Exception:
    colors = None
    A4 = None
    getSampleStyleSheet = None
    Paragraph = None
    SimpleDocTemplate = None
    Spacer = None
    Table = None
    TableStyle = None
    pdfmetrics = None
    TTFont = None

CSV_HEADERS = ["date", "amount", "item", "unit", "payment", "merchant", "category", "excluded"]
ISO_CURRENCIES = sorted({
    "AED","AFN","ALL","AMD","ANG","AOA","ARS","AUD","AWG","AZN","BAM","BBD","BDT","BGN","BHD","BIF",
    "BMD","BND","BOB","BRL","BSD","BTN","BWP","BYN","BZD","CAD","CDF","CHF","CLP","CNY","COP","CRC",
    "CUP","CVE","CZK","DJF","DKK","DOP","DZD","EGP","ERN","ETB","EUR","FJD","FKP","FOK","GBP","GEL",
    "GGP","GHS","GIP","GMD","GNF","GTQ","GYD","HKD","HNL","HRK","HTG","HUF","IDR","ILS","IMP","INR",
    "IQD","IRR","ISK","JEP","JMD","JOD","JPY","KES","KGS","KHR","KID","KMF","KRW","KWD","KYD","KZT",
    "LAK","LBP","LKR","LRD","LSL","LYD","MAD","MDL","MGA","MKD","MMK","MNT","MOP","MRU","MUR","MVR",
    "MWK","MXN","MYR","MZN","NAD","NGN","NIO","NOK","NPR","NZD","OMR","PAB","PEN","PGK","PHP","PKR",
    "PLN","PYG","QAR","RON","RSD","RUB","RWF","SAR","SBD","SCR","SDG","SEK","SGD","SHP","SLE","SLL",
    "SOS","SRD","SSP","STN","SYP","SZL","THB","TJS","TMT","TND","TOP","TRY","TTD","TVD","TWD","TZS",
    "UAH","UGX","USD","UYU","UZS","VED","VES","VND","VUV","WST","XAF","XCD","XCG","XDR","XOF","XPF",
    "YER","ZAR","ZMW","ZWL"
})

I18N = {
    "en": {
        "app_title": "DELR Ledger", "language": "Language", "data_folder": "Data Folder", "choose_folder": "Choose Folder", "current_file": "Current Ledger",
        "new": "New (.delr)", "open": "Open (.delr)", "export": "Export", "import": "Import", "type": "Type", "income": "Income", "expense": "Expense",
        "date": "Date", "item": "Item", "category": "Category", "amount": "Amount", "unit": "Unit", "payment": "Payment", "merchant": "Merchant", "excluded": "Do not count",
        "add": "Add", "clear": "Clear", "edit": "Edit", "delete": "Delete", "entries": "Ledger Entries", "total": "Total", "income_total": "Income", "expense_total": "Expense",
        "filter": "Filter", "from": "From", "to": "To", "invalid_item": "Item cannot be empty.", "invalid_amount": "Amount must be a number.",
        "invalid_date": "Invalid date.", "no_select": "Please select at least one row.", "apply": "Apply", "cancel": "Cancel", "contains": "Contains", "show_all": "All", "show_year": "By Year", "show_month": "By Month", "show_day": "By Day", "prev": "Prev", "next": "Next", "paste_import": "Import from Clipboard", "parse_and_import": "Parse and Import", "import_help": "One record per line. Supported delimiters: Tab, ;, |. The app will detect date/type/amount/unit first, then assign text fields.", "import_date_rule": "Date format: YYYY-MM-DD, YYYY/MM/DD, or YYYY.MM.DD.", "import_summary": "Imported: {ok}, Failed: {bad}", "import_errors": "Import Errors", "choose_field_role": "Choose field role", "as_payment": "As Payment", "as_merchant": "As Merchant"
    },
    "zh-TW": {
        "app_title": "DELR 記賬本", "language": "語言", "data_folder": "資料夾", "choose_folder": "選擇資料夾", "current_file": "目前賬本",
        "new": "新建（.delr）", "open": "開啟（.delr）", "export": "導出", "import": "導入", "type": "類型", "income": "收入", "expense": "支出",
        "date": "日期", "item": "項目", "category": "分類", "amount": "金額", "unit": "單位", "payment": "支付方式", "merchant": "商家", "excluded": "不記入收支",
        "add": "新增", "clear": "清空", "edit": "修改", "delete": "刪除", "entries": "賬本條目", "total": "總計", "income_total": "收入", "expense_total": "支出",
        "filter": "篩選", "from": "起", "to": "迄", "invalid_item": "項目不可為空。", "invalid_amount": "金額必須是數字。",
        "invalid_date": "日期無效。", "no_select": "請至少選擇一行。", "apply": "套用", "cancel": "取消", "contains": "包含", "show_all": "全部", "show_year": "按年顯示", "show_month": "按月顯示", "show_day": "按日顯示", "prev": "上一頁", "next": "下一頁"
    },
    "de": {
        "app_title": "DELR Kassenbuch", "language": "Sprache", "data_folder": "Datenordner", "choose_folder": "Ordner wählen", "current_file": "Aktuelles Ledger",
        "new": "Neu (.delr)", "open": "Öffnen (.delr)", "export": "Exportieren", "import": "Importieren", "type": "Typ", "income": "Einnahme", "expense": "Ausgabe",
        "date": "Datum", "item": "Artikel", "category": "Kategorie", "amount": "Betrag", "unit": "Einheit", "payment": "Zahlung", "merchant": "Händler", "excluded": "Nicht mitzählen",
        "add": "Hinzufügen", "clear": "Leeren", "edit": "Bearbeiten", "delete": "Löschen", "entries": "Ledger-Einträge", "total": "Summe", "income_total": "Einnahmen", "expense_total": "Ausgaben",
        "filter": "Filter", "from": "Von", "to": "Bis", "invalid_item": "Artikel darf nicht leer sein.", "invalid_amount": "Betrag muss eine Zahl sein.",
        "invalid_date": "Ungültiges Datum.", "no_select": "Bitte mindestens eine Zeile wählen.", "apply": "Anwenden", "cancel": "Abbrechen", "contains": "Enthält", "show_all": "Alle", "show_year": "Nach Jahr", "show_month": "Nach Monat", "show_day": "Nach Tag", "prev": "Zurück", "next": "Weiter"
    },
}

LANG_DISPLAY_TO_CODE = {"繁體中文": "zh-TW", "English": "en", "Deutsch": "de"}
CASH_LABELS = {"zh-TW": "現金", "en": "Cash", "de": "Bar"}
ALL_TYPE_LABELS = {
    "income": {"income", "收入", "einnahme"},
    "expense": {"expense", "支出", "ausgabe"},
}


PAYMENT_HINTS = {
    "現金", "cash", "bar", "信用卡", "借記卡", "儲蓄卡", "銀行卡", "刷卡",
    "card", "credit card", "debit card", "karte", "轉賬", "轉帳", "銀行轉賬",
    "銀行轉帳", "bank transfer", "transfer", "überweisung", "ueberweisung",
    "微信支付", "支付寶", "paypal", "apple pay", "google pay"
}


@dataclass
class Entry:
    date: str
    amount: float
    item: str
    unit: str
    payment: str
    merchant: str
    category: str
    excluded: bool = False


class DelrLedgerApp:
    def __init__(self, root: tk.Tk, app_dir: Path) -> None:
        self.root = root
        self.app_dir = app_dir
        self.user_dir = app_dir / "user"
        self.user_dir.mkdir(parents=True, exist_ok=True)
        self.config_dir = app_dir / "config"
        self.config_dir.mkdir(parents=True, exist_ok=True)
        self.settings_path = self.config_dir / "settings.json"

        now = datetime.now()
        self.lang_display_var = tk.StringVar(value="繁體中文")
        self.folder_var = tk.StringVar(value=str(self.user_dir))
        self.file_var = tk.StringVar(value="")
        self.type_var = tk.StringVar(value="expense")
        self.y_var = tk.StringVar(value=str(now.year))
        self.m_var = tk.StringVar(value=f"{now.month:02d}")
        self.d_var = tk.StringVar(value=f"{now.day:02d}")
        self.item_var = tk.StringVar()
        self.category_var = tk.StringVar()
        self.amount_var = tk.StringVar()
        self.unit_var = tk.StringVar(value="")
        self.payment_var = tk.StringVar(value="")
        self.merchant_var = tk.StringVar()
        self.excluded_var = tk.BooleanVar(value=False)

        self.entries: list[Entry] = []
        self.current_file: Path | None = None

        self.header_filters: dict[str, str] = {}
        self.header_range: dict[str, tuple[str, str]] = {}
        self.header_multi: dict[str, set[str]] = {}
        self.sort_column: str | None = None
        self.sort_direction: str | None = None

        self.payment_values: set[str] = set()
        self.category_values: set[str] = set()
        self.merchant_values: set[str] = set()
        self.unit_values: set[str] = set(ISO_CURRENCIES)

        self.view_mode_var = tk.StringVar(value="all")
        self.page_keys: list[str] = []
        self.page_index = 0
        self.doc_current_page_only_var = tk.BooleanVar(value=False)

        self._build_ui()
        self._load_settings()
        self._ensure_clipboard_i18n()
        self.apply_language()
        self._open_last()
        self._update_days()

    def code(self) -> str:
        return LANG_DISPLAY_TO_CODE.get(self.lang_display_var.get(), "zh-TW")

    def tr(self, key: str) -> str:
        lang = self.code()
        if lang in I18N and key in I18N[lang]:
            return I18N[lang][key]
        if key in I18N.get("en", {}):
            return I18N["en"][key]
        return key



    def _ensure_clipboard_i18n(self) -> None:
        # Keep new clipboard-import labels in fallback only, one key exactly once.
        defaults = {
            "en": {
    "paste_import": "Import from Clipboard",
    "parse_and_import": "Parse and Import",
    "import_help": "One record per line. Supported delimiters: Tab, ;, |. The app will detect date/type/amount/unit first, then assign text fields.",
    "import_date_rule": "Date format: YYYY-MM-DD, YYYY/MM/DD, or YYYY.MM.DD.",
    "import_summary": "Imported: {ok}, Failed: {bad}",
    "import_errors": "Import Errors",
    "choose_field_role": "Choose field role",
    "as_payment": "As Payment",
    "as_merchant": "As Merchant",
    "export_doc": "Export as Document",
    "export_current_page_only": "Current page only",
    "doc_title": "Expenses",
    "doc_item": "Item",
    "doc_price": "Price",
    "doc_payment": "Payment",
    "doc_merchant": "Merchant",
    "doc_category": "Category",
    "doc_subtotal": "Subtotal",
    "doc_total": "Total",
    "choose_doc_format": "Choose Document Format",
    "doc_format_md": "Markdown (.md)",
    "doc_format_docx": "Word (.docx)",
    "doc_format_pdf": "PDF (.pdf)"
},
"zh-TW": {
    "paste_import": "從剪貼板匯入",
    "parse_and_import": "解析並匯入",
    "import_help": "每行一筆資料；支援 Tab、;、|。程式會先識別日期、類型、金額、單位，其餘文本再分配到項目、支付方式、商家、分類。",
    "import_date_rule": "日期格式：YYYY-MM-DD、YYYY/MM/DD 或 YYYY.MM.DD。",
    "import_summary": "成功匯入 {ok} 筆，失敗 {bad} 筆",
    "import_errors": "匯入錯誤",
    "choose_field_role": "選擇欄位角色",
    "as_payment": "支付方式",
    "as_merchant": "商家",
    "export_doc": "導出爲文檔",
    "export_current_page_only": "僅導出當前頁",
    "doc_title": "支出記錄",
    "doc_item": "商品名",
    "doc_price": "價格",
    "doc_payment": "支付方式",
    "doc_merchant": "消費地點",
    "doc_category": "分類",
    "doc_subtotal": "小計",
    "doc_total": "總計",
    "choose_doc_format": "選擇文檔格式",
    "doc_format_md": "Markdown (.md)",
    "doc_format_docx": "Word (.docx)",
    "doc_format_pdf": "PDF (.pdf)"
},
"de": {
    "paste_import": "Aus Zwischenablage importieren",
    "parse_and_import": "Analysieren und importieren",
    "import_help": "Eine Zeile pro Datensatz; Trennzeichen: Tab, ;, |. Das Programm erkennt zuerst Datum, Typ, Betrag und Einheit.",
    "import_date_rule": "Datumsformat: JJJJ-MM-TT, JJJJ/MM/TT oder JJJJ.MM.TT.",
    "import_summary": "Erfolgreich: {ok}, Fehlgeschlagen: {bad}",
    "import_errors": "Importfehler",
    "choose_field_role": "Feldrolle wählen",
    "as_payment": "Zahlungsart",
    "as_merchant": "Händler",
    "export_doc": "Als Dokument exportieren",
    "export_current_page_only": "Nur aktuelle Seite",
    "doc_title": "Ausgaben",
    "doc_item": "Artikel",
    "doc_price": "Preis",
    "doc_payment": "Zahlungsart",
    "doc_merchant": "Ort",
    "doc_category": "Kategorie",
    "doc_subtotal": "Zwischensumme",
    "doc_total": "Gesamtsumme",
    "choose_doc_format": "Dokumentformat wählen",
    "doc_format_md": "Markdown (.md)",
    "doc_format_docx": "Word (.docx)",
    "doc_format_pdf": "PDF (.pdf)"
}
        }
        for lang, kv in defaults.items():
            I18N.setdefault(lang, {})
            for k, v in kv.items():
                I18N[lang].setdefault(k, v)

    def cash(self) -> str:
        return CASH_LABELS[self.code()]

    def ui_date_fmt(self) -> str:
        return "%m-%d-%Y" if self.code() == "en" else ("%d-%m-%Y" if self.code() == "de" else "%Y-%m-%d")

    def fmt_ui_date(self, iso: str) -> str:
        return datetime.strptime(iso, "%Y-%m-%d").strftime(self.ui_date_fmt())
    def _build_ui(self) -> None:
        self.root.title("DELR Ledger")
        self.root.geometry("1280x760")
        self.root.minsize(1120, 620)

        top = ttk.Frame(self.root)
        top.pack(fill="x", padx=10, pady=(10, 6))
        self.view_all_rb = ttk.Radiobutton(top, variable=self.view_mode_var, value="all", command=self._on_view_mode_changed)
        self.view_all_rb.pack(side="left")
        self.view_year_rb = ttk.Radiobutton(top, variable=self.view_mode_var, value="year", command=self._on_view_mode_changed)
        self.view_year_rb.pack(side="left", padx=(6, 0))
        self.view_month_rb = ttk.Radiobutton(top, variable=self.view_mode_var, value="month", command=self._on_view_mode_changed)
        self.view_month_rb.pack(side="left", padx=(6, 0))
        self.view_day_rb = ttk.Radiobutton(top, variable=self.view_mode_var, value="day", command=self._on_view_mode_changed)
        self.view_day_rb.pack(side="left", padx=(6, 12))
        self.lang_label = ttk.Label(top)
        self.lang_label.pack(side="left")
        self.lang_combo = ttk.Combobox(top, textvariable=self.lang_display_var, values=["繁體中文", "English", "Deutsch"], state="readonly", width=12)
        self.lang_combo.pack(side="left", padx=(6, 14))
        self.lang_combo.bind("<<ComboboxSelected>>", lambda _e: self.on_lang())
        self.folder_label = ttk.Label(top)
        self.folder_label.pack(side="left")
        ttk.Entry(top, textvariable=self.folder_var, width=42).pack(side="left", padx=6)
        self.choose_btn = ttk.Button(top, command=self.choose_folder)
        self.choose_btn.pack(side="left")

        file_row = ttk.Frame(self.root)
        file_row.pack(fill="x", padx=10, pady=(0, 6))
        self.file_label = ttk.Label(file_row)
        self.file_label.pack(side="left")
        ttk.Entry(file_row, textvariable=self.file_var).pack(side="left", fill="x", expand=True, padx=6)
        self.new_btn = ttk.Button(file_row, command=self.new_ledger)
        self.new_btn.pack(side="left")
        self.open_btn = ttk.Button(file_row, command=self.open_ledger)
        self.open_btn.pack(side="left", padx=(6, 0))
        self.export_btn = ttk.Button(file_row, command=self.export_ledger)
        self.export_btn.pack(side="left", padx=(6, 0))
        self.import_btn = ttk.Button(file_row, command=self.import_merge)
        self.import_btn.pack(side="left", padx=(6, 0))

        self.export_doc_btn = ttk.Button(file_row, command=self.export_document)
        self.export_doc_btn.pack(side="right")
        self.doc_current_page_only_chk = ttk.Checkbutton(file_row, variable=self.doc_current_page_only_var)
        self.doc_current_page_only_chk.pack(side="right", padx=(0, 8))

        frm = ttk.LabelFrame(self.root)
        frm.pack(fill="x", padx=10, pady=(0, 6))
        self.form_frame = frm

        r1 = ttk.Frame(frm)
        r1.pack(fill="x", padx=8, pady=(6, 2))
        self.type_label = ttk.Label(r1)
        self.type_label.pack(side="left")
        self.type_combo = ttk.Combobox(r1, textvariable=self.type_var, state="readonly", width=10)
        self.type_combo.pack(side="left", padx=(4, 14))

        self.date_label = ttk.Label(r1)
        self.date_label.pack(side="left")
        self.date_wrap = ttk.Frame(r1)
        self.date_wrap.pack(side="left", padx=(4, 14))
        self.y_combo = ttk.Combobox(self.date_wrap, textvariable=self.y_var, values=[str(y) for y in range(2000, 2201)], width=6, state="readonly")
        self.y_u = ttk.Label(self.date_wrap)
        self.m_combo = ttk.Combobox(self.date_wrap, textvariable=self.m_var, values=[f"{m:02d}" for m in range(1, 13)], width=4, state="readonly")
        self.m_u = ttk.Label(self.date_wrap)
        self.d_combo = ttk.Combobox(self.date_wrap, textvariable=self.d_var, values=[f"{d:02d}" for d in range(1, 32)], width=4, state="readonly")
        self.d_u = ttk.Label(self.date_wrap)
        self.y_combo.bind("<<ComboboxSelected>>", lambda _e: self._update_days())
        self.m_combo.bind("<<ComboboxSelected>>", lambda _e: self._update_days())

        self.item_label = ttk.Label(r1)
        self.item_label.pack(side="left")
        ttk.Entry(r1, textvariable=self.item_var, width=20).pack(side="left", padx=(4, 14))

        self.category_label = ttk.Label(r1)
        self.category_label.pack(side="left")
        self.category_combo = ttk.Combobox(r1, textvariable=self.category_var, values=[], width=16)
        self.category_combo.pack(side="left", padx=(4, 8))

        r2 = ttk.Frame(frm)
        r2.pack(fill="x", padx=8, pady=(2, 6))
        self.amount_label = ttk.Label(r2)
        self.amount_label.pack(side="left")
        ttk.Entry(r2, textvariable=self.amount_var, width=12).pack(side="left", padx=(4, 14))

        self.unit_label = ttk.Label(r2)
        self.unit_label.pack(side="left")
        self.unit_combo = ttk.Combobox(r2, textvariable=self.unit_var, values=sorted(self.unit_values), width=8)
        self.unit_combo.pack(side="left", padx=(4, 14))
        self.unit_combo.bind("<KeyRelease>", lambda _e: self._filter_combo_values(self.unit_combo, sorted(self.unit_values)))
        self.unit_combo.bind("<<ComboboxSelected>>", lambda _e: self._save_settings())
        self.unit_combo.bind("<FocusOut>", lambda _e: self._save_settings())

        self.payment_label = ttk.Label(r2)
        self.payment_label.pack(side="left")
        self.payment_combo = ttk.Combobox(r2, textvariable=self.payment_var, values=[], width=16)
        self.payment_combo.pack(side="left", padx=(4, 14))

        self.merchant_label = ttk.Label(r2)
        self.merchant_label.pack(side="left")
        self.merchant_combo = ttk.Combobox(r2, textvariable=self.merchant_var, values=[], width=22)
        self.merchant_combo.pack(side="left", padx=(4, 8))

        self.excluded_chk = ttk.Checkbutton(r2, variable=self.excluded_var)
        self.excluded_chk.pack(side="left", padx=(0, 8))

        ttk.Frame(r2).pack(side="left", fill="x", expand=True)
        self.add_btn = ttk.Button(r2, command=self.add_entry)
        self.add_btn.pack(side="right")
        self.clear_btn = ttk.Button(r2, command=self.clear_form)
        self.clear_btn.pack(side="right", padx=(0, 6))
        self.paste_import_btn = ttk.Button(r2, command=self.open_paste_import_window)
        self.paste_import_btn.pack(side="right", padx=(0, 6))

        table_box = ttk.LabelFrame(self.root)
        table_box.pack(fill="both", expand=True, padx=10, pady=(0, 6))
        self.table_box = table_box
        hdr = ttk.Frame(table_box)
        hdr.pack(fill="x", padx=8, pady=(6, 4))
        self.page_prev_btn = ttk.Button(hdr, command=lambda: self._change_page(-1))
        self.page_prev_btn.pack(side="left")
        self.page_next_btn = ttk.Button(hdr, command=lambda: self._change_page(1))
        self.page_next_btn.pack(side="left", padx=(6, 0))
        self.page_info_label = ttk.Label(hdr, width=18)
        self.page_info_label.pack(side="left", padx=(8, 6))
        ttk.Frame(hdr).pack(side="left", fill="x", expand=True)
        self.del_btn = ttk.Button(hdr, command=self.delete_selected, state="disabled")
        self.del_btn.pack(side="right")
        self.edit_btn = ttk.Button(hdr, command=self.edit_selected, state="disabled")
        self.edit_btn.pack(side="right", padx=(0, 6))

        tw = ttk.Frame(table_box)
        tw.pack(fill="both", expand=True, padx=8, pady=(0, 8))
        cols = ("date", "type", "item", "amount", "unit", "payment", "merchant", "category", "excluded")
        self.tree = ttk.Treeview(tw, columns=cols, show="headings", selectmode="extended", height=8)
        widths = {"date": 120, "type": 90, "item": 220, "amount": 100, "unit": 70, "payment": 130, "merchant": 180, "category": 130, "excluded": 120}
        for c in cols:
            self.tree.column(c, width=widths[c], anchor=("e" if c == "amount" else "w"))
            self.tree.heading(c, text=c, command=lambda col=c: self._on_header_left_click(col))
        y_scroll = ttk.Scrollbar(tw, orient="vertical", command=self.tree.yview)
        self.tree.configure(yscrollcommand=y_scroll.set)
        self.tree.pack(side="left", fill="both", expand=True)
        y_scroll.pack(side="left", fill="y")
        self.tree.bind("<<TreeviewSelect>>", lambda _e: self._update_action_buttons())
        self.tree.bind("<Button-3>", self._on_treeview_right_click)

        st = ttk.Frame(self.root)
        st.pack(side="bottom", fill="x", padx=10, pady=(0, 10))
        self.total_label = ttk.Label(st)
        self.total_label.pack(anchor="w")
        self.income_label = ttk.Label(st)
        self.income_label.pack(anchor="w")
        self.expense_label = ttk.Label(st)
        self.expense_label.pack(anchor="w")

    def _layout_date(self) -> None:
        for child in self.date_wrap.winfo_children():
            child.pack_forget()
        if self.code() == "en":
            items = [(self.m_combo, "M"), (self.d_combo, "D"), (self.y_combo, "Y")]
        elif self.code() == "de":
            items = [(self.d_combo, "T"), (self.m_combo, "M"), (self.y_combo, "J")]
        else:
            items = [(self.y_combo, "年"), (self.m_combo, "月"), (self.d_combo, "日")]
        labels = [self.y_u, self.m_u, self.d_u]
        for (combo, text), lbl in zip(items, labels):
            combo.pack(side="left", padx=(0, 1))
            lbl.config(text=text)
            lbl.pack(side="left", padx=(0, 2))

    def _load_settings(self) -> None:
        if not self.settings_path.exists():
            return
        try:
            settings = json.loads(self.settings_path.read_text(encoding="utf-8"))
        except Exception:
            return
        if settings.get("lang_display") in LANG_DISPLAY_TO_CODE:
            self.lang_display_var.set(settings["lang_display"])
        if isinstance(settings.get("data_folder"), str) and settings["data_folder"].strip():
            self.folder_var.set(settings["data_folder"])
        if isinstance(settings.get("last_file"), str) and settings["last_file"].strip():
            self.file_var.set(settings["last_file"])
        if isinstance(settings.get("last_unit"), str) and settings["last_unit"].strip().upper() in ISO_CURRENCIES:
            self.unit_var.set(settings["last_unit"].strip().upper())
        if isinstance(settings.get("last_payment"), str):
            self.payment_var.set(settings["last_payment"].strip())

    def _save_settings(self) -> None:
        payload = {"lang_display": self.lang_display_var.get(), "data_folder": self.folder_var.get(), "last_file": str(self.current_file) if self.current_file else "", "last_unit": self.unit_var.get().strip().upper(), "last_payment": self.payment_var.get().strip()}
        self.settings_path.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")

    def _type_from_ui(self, value: str) -> str:
        s = (value or "").strip().lower()
        if s in ALL_TYPE_LABELS["income"]:
            return "income"
        return "expense"

    def _filter_combo_values(self, combo: ttk.Combobox, source: list[str]) -> None:
        typed = combo.get()
        key = typed.strip().upper()
        matched = source if not key else [v for v in source if key in v.upper()]
        combo["values"] = matched
        if matched and combo.focus_get() == combo:
            combo.event_generate("<Down>")
            combo.after_idle(lambda c=combo, t=typed: c.set(t))

    def _enforce_combo_legal(self, combo: ttk.Combobox, source: list[str], fallback: str = "") -> None:
        value = combo.get().strip().upper()
        legal = {x.upper(): x for x in source}
        combo.set(legal[value] if value in legal else fallback)

    def on_lang(self) -> None:
        self.apply_language()
        self.refresh_table()
        self._save_settings()

    def _update_action_buttons(self) -> None:
        selected = len(self.tree.selection())
        self.del_btn.config(state=("normal" if selected > 0 else "disabled"))
        self.edit_btn.config(state=("normal" if selected == 1 else "disabled"))

    def _on_view_mode_changed(self) -> None:
        self.page_index = 0
        self.refresh_table()

    def _make_page_key(self, date_iso: str) -> str:
        mode = self.view_mode_var.get()
        if mode == "year":
            return date_iso[:4]
        if mode == "month":
            return date_iso[:7]
        if mode == "day":
            return date_iso[:10]
        return "all"

    def _change_page(self, delta: int) -> None:
        if not self.page_keys:
            return
        new_idx = self.page_index + delta
        if 0 <= new_idx < len(self.page_keys):
            self.page_index = new_idx
            self.refresh_table()

    def _on_header_left_click(self, col: str) -> None:
        if self.sort_column != col:
            self.sort_column = col
            self.sort_direction = "asc"
        elif self.sort_direction == "asc":
            self.sort_direction = "desc"
        else:
            self.sort_column = None
            self.sort_direction = None
        self.refresh_table()

    def _column_from_tree_ident(self, ident: str) -> str | None:
        if not ident or not ident.startswith("#"):
            return None
        cols = ("date", "type", "item", "amount", "unit", "payment", "merchant", "category", "excluded")
        try:
            idx = int(ident[1:]) - 1
        except ValueError:
            return None
        if 0 <= idx < len(cols):
            return cols[idx]
        return None

    def _on_treeview_right_click(self, event: tk.Event) -> None:
        if self.tree.identify_region(event.x, event.y) != "heading":
            return
        col_ident = self.tree.identify_column(event.x)
        col = self._column_from_tree_ident(col_ident)
        if col:
            self.open_filter_for_column(col)

    def open_filter_for_column(self, col: str) -> None:
        if col == "date":
            self._toggle_or_apply_filter(col, lambda: self._open_date_filter(col))
        elif col == "type":
            self._toggle_or_apply_filter(col, lambda: self._open_type_filter(col))
        elif col == "item":
            self._toggle_or_apply_filter(col, lambda: self._open_text_filter(col))
        elif col in {"payment", "merchant", "category"}:
            self._toggle_or_apply_filter(col, lambda c=col: self._open_multi_filter(c))
        elif col in {"amount", "unit"}:
            self._toggle_or_apply_filter("amount", self._open_amount_unit_filter)
        elif col == "excluded":
            self._toggle_or_apply_filter(col, lambda: self._apply_excluded_filter(col))
        else:
            self._toggle_or_apply_filter(col, lambda: self._open_text_filter(col))

    def _update_paging_controls(self) -> None:
        mode = self.view_mode_var.get()
        if mode == "all" or not self.page_keys:
            self.page_prev_btn.config(state="disabled")
            self.page_next_btn.config(state="disabled")
            self.page_info_label.config(text="")
            return
        current = self.page_keys[self.page_index]
        self.page_prev_btn.config(state=("normal" if self.page_index > 0 else "disabled"))
        self.page_next_btn.config(state=("normal" if self.page_index < len(self.page_keys) - 1 else "disabled"))
        self.page_info_label.config(text=f"{self.page_index + 1}/{len(self.page_keys)}  {current}")


    def _toggle_or_apply_filter(self, col: str, apply_fn) -> None:
        has_filter = col in self.header_filters or col in self.header_range or col in self.header_multi
        if has_filter:
            self.header_filters.pop(col, None)
            self.header_range.pop(col, None)
            self.header_multi.pop(col, None)
            if col in {"amount", "unit"}:
                self.header_filters.pop("unit", None)
                self.header_range.pop("amount", None)
            self.refresh_table()
            return
        apply_fn()

    def _open_filter_window(self, title: str) -> tk.Toplevel:
        w = tk.Toplevel(self.root)
        w.title(title)
        w.resizable(False, False)
        w.transient(self.root)
        w.grab_set()
        return w

    def _finalize_filter_window(self, w: tk.Toplevel, pad_w: int = 20, pad_h: int = 20) -> None:
        self.root.update_idletasks(); w.update_idletasks()
        req_w = w.winfo_reqwidth() + pad_w; req_h = w.winfo_reqheight() + pad_h
        x = self.root.winfo_rootx() + max((self.root.winfo_width() - req_w)//2, 0)
        y = self.root.winfo_rooty() + max((self.root.winfo_height() - req_h)//2, 0)
        w.geometry(f"{req_w}x{req_h}+{x}+{y}")

    def _open_date_filter(self, col: str) -> None:
        w = self._open_filter_window(self.tr("filter"))
        now = datetime.now()
        fy, fm, fd = tk.StringVar(value=str(now.year)), tk.StringVar(value=f"{now.month:02d}"), tk.StringVar(value=f"{now.day:02d}")
        ty, tm, td = tk.StringVar(value=str(now.year)), tk.StringVar(value=f"{now.month:02d}"), tk.StringVar(value=f"{now.day:02d}")
        r1 = ttk.Frame(w); r1.pack(fill="x", padx=10, pady=(12, 6)); ttk.Label(r1, text=f"{self.tr('from')}:").pack(side="left"); self._create_date_picker(r1, (fy, fm, fd))
        r2 = ttk.Frame(w); r2.pack(fill="x", padx=10, pady=(0, 10)); ttk.Label(r2, text=f"{self.tr('to')}:").pack(side="left"); self._create_date_picker(r2, (ty, tm, td))
        btn = ttk.Frame(w); btn.pack(side="bottom", fill="x", padx=10, pady=10)
        def apply() -> None:
            self.header_range[col] = (f"{fy.get()}-{fm.get()}-{fd.get()}", f"{ty.get()}-{tm.get()}-{td.get()}")
            w.destroy(); self.refresh_table()
        ttk.Button(btn, text=self.tr("apply"), command=apply).pack(side="right"); ttk.Button(btn, text=self.tr("cancel"), command=w.destroy).pack(side="right", padx=(0, 6)); self._finalize_filter_window(w)

    def _open_type_filter(self, col: str) -> None:
        w = self._open_filter_window(self.tr("filter")); values = [self.tr("expense"), self.tr("income")]; var = tk.StringVar(value=values[0])
        row = ttk.Frame(w); row.pack(fill="x", padx=10, pady=14); ttk.Label(row, text=f"{self.tr('type')}:").pack(side="left"); ttk.Combobox(row, textvariable=var, values=values, state="readonly", width=20).pack(side="left", padx=(6, 0))
        btn = ttk.Frame(w); btn.pack(side="bottom", fill="x", padx=10, pady=10)
        def apply() -> None:
            self.header_filters[col] = self._type_from_ui(var.get()); w.destroy(); self.refresh_table()
        ttk.Button(btn, text=self.tr("apply"), command=apply).pack(side="right"); ttk.Button(btn, text=self.tr("cancel"), command=w.destroy).pack(side="right", padx=(0, 6)); self._finalize_filter_window(w)

    def _open_text_filter(self, col: str) -> None:
        w = self._open_filter_window(self.tr("filter")); var = tk.StringVar(value="")
        row = ttk.Frame(w); row.pack(fill="x", padx=10, pady=14); ttk.Label(row, text=f"{self.tr('contains')}: ").pack(side="left"); ttk.Entry(row, textvariable=var, width=40).pack(side="left", padx=(6, 0))
        btn = ttk.Frame(w); btn.pack(side="bottom", fill="x", padx=10, pady=10)
        def apply() -> None:
            value = var.get().strip();
            if value: self.header_filters[col] = value
            w.destroy(); self.refresh_table()
        ttk.Button(btn, text=self.tr("apply"), command=apply).pack(side="right"); ttk.Button(btn, text=self.tr("cancel"), command=w.destroy).pack(side="right", padx=(0, 6)); self._finalize_filter_window(w)

    def _open_multi_filter(self, col: str) -> None:
        w = self._open_filter_window(self.tr("filter")); values = sorted({"payment": self.payment_values, "merchant": self.merchant_values, "category": self.category_values}[col])
        row = ttk.Frame(w); row.pack(fill="both", expand=True, padx=10, pady=10)
        listbox = tk.Listbox(row, selectmode="multiple", exportselection=False); listbox.pack(side="left", fill="both", expand=True)
        scroll = ttk.Scrollbar(row, orient="vertical", command=listbox.yview); listbox.configure(yscrollcommand=scroll.set); scroll.pack(side="left", fill="y")
        for v in values: listbox.insert("end", v)
        btn = ttk.Frame(w); btn.pack(side="bottom", fill="x", padx=10, pady=10)
        def apply() -> None:
            picks = {listbox.get(i) for i in listbox.curselection()}
            if picks: self.header_multi[col] = picks
            w.destroy(); self.refresh_table()
        ttk.Button(btn, text=self.tr("apply"), command=apply).pack(side="right"); ttk.Button(btn, text=self.tr("cancel"), command=w.destroy).pack(side="right", padx=(0, 6)); self._finalize_filter_window(w, pad_w=30, pad_h=30)

    def _open_amount_unit_filter(self) -> None:
        w = self._open_filter_window(self.tr("filter")); unit_var = tk.StringVar(value="EUR"); lo_var = tk.StringVar(value=""); hi_var = tk.StringVar(value="")
        r1 = ttk.Frame(w); r1.pack(fill="x", padx=10, pady=(12, 6)); ttk.Label(r1, text=f"{self.tr('unit')}:").pack(side="left")
        u_combo = ttk.Combobox(r1, textvariable=unit_var, values=sorted(self.unit_values), width=10); u_combo.pack(side="left", padx=(6, 0)); u_combo.bind("<KeyRelease>", lambda _e: self._filter_combo_values(u_combo, sorted(self.unit_values))); u_combo.bind("<FocusOut>", lambda _e: None)
        r2 = ttk.Frame(w); r2.pack(fill="x", padx=10, pady=(0, 6)); ttk.Label(r2, text=f"{self.tr('from')}:").pack(side="left"); ttk.Entry(r2, textvariable=lo_var, width=14).pack(side="left", padx=(6, 14)); ttk.Label(r2, text=f"{self.tr('to')}:").pack(side="left"); ttk.Entry(r2, textvariable=hi_var, width=14).pack(side="left", padx=(6, 0))
        btn = ttk.Frame(w); btn.pack(side="bottom", fill="x", padx=10, pady=10)
        def apply() -> None:
            unit = unit_var.get().strip().upper()
            if unit not in self.unit_values: return
            self.header_filters["unit"] = unit; self.header_range["amount"] = (lo_var.get().strip(), hi_var.get().strip()); w.destroy(); self.refresh_table()
        ttk.Button(btn, text=self.tr("apply"), command=apply).pack(side="right"); ttk.Button(btn, text=self.tr("cancel"), command=w.destroy).pack(side="right", padx=(0, 6)); self._finalize_filter_window(w)

    def _apply_excluded_filter(self, col: str) -> None:
        self.header_filters[col] = "1"
        self.refresh_table()

    def _filtered(self) -> list[tuple[int, Entry]]:
        out: list[tuple[int, Entry]] = []
        for i, e in enumerate(self.entries):
            if not self._passes_filters(e):
                continue
            out.append((i, e))
        return out

    def _passes_filters(self, e: Entry) -> bool:
        ftype = self.header_filters.get("type", "")
        if ftype == "income" and e.amount >= 0:
            return False
        if ftype == "expense" and e.amount <= 0:
            return False
        if self.header_filters.get("item") and self.header_filters["item"].lower() not in e.item.lower():
            return False
        if self.header_filters.get("unit") and self.header_filters["unit"].upper() != e.unit.upper():
            return False
        if self.header_filters.get("excluded") == "1" and not e.excluded:
            return False
        if "payment" in self.header_multi and e.payment not in self.header_multi["payment"]:
            return False
        if "merchant" in self.header_multi and e.merchant not in self.header_multi["merchant"]:
            return False
        if "category" in self.header_multi and e.category not in self.header_multi["category"]:
            return False

        aa = abs(e.amount)
        if "amount" in self.header_range:
            lo, hi = self.header_range["amount"]
            try:
                if lo and aa < float(lo):
                    return False
            except ValueError:
                pass
            try:
                if hi and aa > float(hi):
                    return False
            except ValueError:
                pass

        if "date" in self.header_range:
            d1, d2 = self.header_range["date"]
            if d1 and e.date < d1:
                return False
            if d2 and e.date > d2:
                return False
        return True

    def _is_filter_active_for_column(self, col: str) -> bool:
        if col in {"amount", "unit"}:
            return ("amount" in self.header_range) or ("unit" in self.header_filters)
        return (col in self.header_filters) or (col in self.header_range) or (col in self.header_multi)

    def update_header_state(self) -> None:
        for col, key in [("date", "date"), ("type", "type"), ("item", "item"), ("amount", "amount"), ("unit", "unit"), ("payment", "payment"), ("merchant", "merchant"), ("category", "category"), ("excluded", "excluded")]:
            label = self.tr(key)
            if self._is_filter_active_for_column(col):
                label = f"{label} [F]"
            if self.sort_column == col and self.sort_direction == "asc":
                label = f"{label} ↑"
            elif self.sort_column == col and self.sort_direction == "desc":
                label = f"{label} ↓"
            self.tree.heading(col, text=label, command=lambda c=col: self._on_header_left_click(c))

    def apply_language(self) -> None:
        self.root.title(self.tr("app_title"))
        self.lang_label.config(text=f"{self.tr('language')}: ")
        self.folder_label.config(text=f"{self.tr('data_folder')}: ")
        self.choose_btn.config(text=self.tr("choose_folder"))
        self.file_label.config(text=f"{self.tr('current_file')}: ")
        self.new_btn.config(text=self.tr("new"))
        self.open_btn.config(text=self.tr("open"))
        self.export_btn.config(text=self.tr("export"))
        self.import_btn.config(text=self.tr("import"))
        self.export_doc_btn.config(text=self.tr("export_doc"))
        self.doc_current_page_only_chk.config(text=self.tr("export_current_page_only"))
        self.form_frame.config(text=self.tr("entries"))
        self.type_label.config(text=self.tr("type")); self.date_label.config(text=self.tr("date")); self.item_label.config(text=self.tr("item")); self.category_label.config(text=self.tr("category"))
        self.amount_label.config(text=self.tr("amount")); self.unit_label.config(text=self.tr("unit")); self.payment_label.config(text=self.tr("payment")); self.merchant_label.config(text=self.tr("merchant")); self.excluded_chk.config(text=self.tr("excluded"))
        self.view_all_rb.config(text=self.tr("show_all"))
        self.view_year_rb.config(text=self.tr("show_year"))
        self.view_month_rb.config(text=self.tr("show_month"))
        self.view_day_rb.config(text=self.tr("show_day"))
        self.page_prev_btn.config(text=self.tr("prev"))
        self.page_next_btn.config(text=self.tr("next"))
        self.add_btn.config(text=self.tr("add")); self.clear_btn.config(text=self.tr("clear")); self.paste_import_btn.config(text=self.tr("paste_import")); self.edit_btn.config(text=self.tr("edit")); self.del_btn.config(text=self.tr("delete"))
        current_type = self._type_from_ui(self.type_var.get())
        self.type_combo["values"] = [self.tr("expense"), self.tr("income")]
        self.type_var.set(self.tr("income") if current_type == "income" else self.tr("expense"))
        current_payment = self.payment_var.get().strip()
        if current_payment and current_payment in set(CASH_LABELS.values()):
            self.payment_var.set(self.cash())
        self._layout_date()
        self.update_header_state()
        self._rebuild_options()

    def choose_folder(self) -> None:
        picked = filedialog.askdirectory(initialdir=self.folder_var.get())
        if picked:
            self.folder_var.set(picked)
            self._save_settings()

    def _ensure_current(self) -> None:
        if self.current_file is not None:
            return
        folder = Path(self.folder_var.get().strip() or str(self.user_dir))
        folder.mkdir(parents=True, exist_ok=True)
        self.current_file = folder / datetime.now().strftime("ledger-%Y%m%d-%H%M%S.delr")
        self.file_var.set(str(self.current_file))
    def _open_last(self) -> None:
        path = self.file_var.get().strip()
        if path and Path(path).exists() and Path(path).suffix.lower() == ".delr":
            self.current_file = Path(path)
            self.entries = self.read_entries_from_path(self.current_file)
            self._rebuild_options()
            self.refresh_table()

    def _read_rows_csv_like(self, path: Path, delimiter: str = ",") -> list[dict[str, str]]:
        rows: list[dict[str, str]] = []
        with path.open("r", encoding="utf-8-sig", newline="") as f:
            for row in csv.DictReader(f, delimiter=delimiter):
                rows.append({k: (row.get(k, "") or "") for k in CSV_HEADERS})
        return rows

    def _read_rows_json(self, path: Path) -> list[dict[str, str]]:
        data = json.loads(path.read_text(encoding="utf-8"))
        if not isinstance(data, list):
            return []
        rows: list[dict[str, str]] = []
        for it in data:
            if isinstance(it, dict):
                rows.append({k: str(it.get(k, "") or "") for k in CSV_HEADERS})
        return rows

    def _read_rows_xml(self, path: Path) -> list[dict[str, str]]:
        root = xml_parse(path).getroot()
        rows: list[dict[str, str]] = []
        for node in root.findall("entry"):
            rows.append({k: (node.findtext(k, default="") or "") for k in CSV_HEADERS})
        return rows

    def _read_rows_yaml(self, path: Path) -> list[dict[str, str]]:
        if yaml is None:
            raise RuntimeError("PyYAML is not installed")
        data = yaml.safe_load(path.read_text(encoding="utf-8"))
        if not isinstance(data, list):
            return []
        rows: list[dict[str, str]] = []
        for it in data:
            if isinstance(it, dict):
                rows.append({k: str(it.get(k, "") or "") for k in CSV_HEADERS})
        return rows

    def _read_rows_xlsx(self, path: Path) -> list[dict[str, str]]:
        if load_workbook is None:
            raise RuntimeError("openpyxl is not installed")
        wb = load_workbook(filename=path, data_only=True)
        ws = wb.active
        header = [str(c.value or "").strip() for c in ws[1]]
        idx = {h: i for i, h in enumerate(header)}
        rows: list[dict[str, str]] = []
        for r in ws.iter_rows(min_row=2):
            row = {}
            for k in CSV_HEADERS:
                i = idx.get(k)
                row[k] = "" if i is None else str(r[i].value or "")
            rows.append(row)
        return rows

    def _rows_to_entries(self, rows: list[dict[str, str]]) -> list[Entry]:
        out: list[Entry] = []
        for r in rows:
            try:
                amount = float(str(r.get("amount", "0") or "0"))
            except ValueError:
                amount = 0.0
            excluded = str(r.get("excluded", "") or "").strip().casefold() in {"1", "true", "yes", "y", "on"}
            out.append(Entry(date=str(r.get("date", "1970-01-01") or "1970-01-01"), amount=amount, item=str(r.get("item", "") or ""), unit=str(r.get("unit", "EUR") or "EUR").upper(), payment=str(r.get("payment", "") or ""), merchant=str(r.get("merchant", "") or ""), category=str(r.get("category", "") or ""), excluded=excluded))
        return out

    def _sorted(self, rows: list[tuple[int, Entry]]) -> list[tuple[int, Entry]]:
        if self.sort_column is None or self.sort_direction is None:
            merchant_order: dict[tuple[str, str], int] = {}
            for idx, e in enumerate(self.entries):
                key = (e.date, (e.merchant or "").casefold())
                merchant_order.setdefault(key, idx)
            return sorted(rows, key=lambda pair: (pair[1].date, merchant_order.get((pair[1].date, (pair[1].merchant or "").casefold()), pair[0]), pair[0]))

        reverse = self.sort_direction == "desc"

        def key_func(pair: tuple[int, Entry]):
            _idx, e = pair
            if self.sort_column == "date":
                return e.date
            if self.sort_column == "type":
                if e.amount < 0:
                    return -1
                if e.amount > 0:
                    return 1
                return 0
            if self.sort_column == "item":
                return e.item.lower()
            if self.sort_column == "amount":
                return abs(e.amount)
            if self.sort_column == "unit":
                return (e.unit or "").upper()
            if self.sort_column == "payment":
                return (e.payment or "").lower()
            if self.sort_column == "merchant":
                return (e.merchant or "").lower()
            if self.sort_column == "category":
                return (e.category or "").lower()
            if self.sort_column == "excluded":
                return 1 if e.excluded else 0
            return ""

        return sorted(rows, key=key_func, reverse=reverse)

    def edit_selected(self) -> None:
        ids = list(self.tree.selection())
        if len(ids) != 1:
            return
        idx = int(ids[0])
        if idx < 0 or idx >= len(self.entries):
            return
        e = self.entries[idx]

        w = self._open_filter_window(self.tr("edit"))

        type_var = tk.StringVar(value=self.tr("income") if e.amount < 0 else self.tr("expense"))
        date_var = tk.StringVar(value=self.fmt_ui_date(e.date))
        item_var = tk.StringVar(value=e.item)
        amount_var = tk.StringVar(value=f"{abs(e.amount):.2f}")
        unit_var = tk.StringVar(value=(e.unit or ""))
        payment_var = tk.StringVar(value=(e.payment if e.payment != "-" else ""))
        merchant_var = tk.StringVar(value=(e.merchant if e.merchant != "-" else ""))
        category_var = tk.StringVar(value=(e.category if e.category != "-" else ""))
        excluded_var = tk.BooleanVar(value=e.excluded)

        r1 = ttk.Frame(w); r1.pack(fill="x", padx=10, pady=(12, 6))
        ttk.Label(r1, text=f"{self.tr('type')}: ").pack(side="left")
        ttk.Combobox(r1, textvariable=type_var, values=[self.tr("expense"), self.tr("income")], state="readonly", width=10).pack(side="left", padx=(4, 10))
        ttk.Label(r1, text=f"{self.tr('date')}: ").pack(side="left")
        ttk.Entry(r1, textvariable=date_var, width=12).pack(side="left", padx=(4, 10))
        ttk.Label(r1, text=f"{self.tr('item')}: ").pack(side="left")
        ttk.Entry(r1, textvariable=item_var, width=22).pack(side="left", padx=(4, 0))

        r2 = ttk.Frame(w); r2.pack(fill="x", padx=10, pady=(0, 6))
        ttk.Label(r2, text=f"{self.tr('amount')}: ").pack(side="left")
        ttk.Entry(r2, textvariable=amount_var, width=12).pack(side="left", padx=(4, 10))
        ttk.Label(r2, text=f"{self.tr('unit')}: ").pack(side="left")
        u_combo = ttk.Combobox(r2, textvariable=unit_var, values=sorted(self.unit_values), width=8)
        u_combo.pack(side="left", padx=(4, 10))
        u_combo.bind("<KeyRelease>", lambda _e: self._filter_combo_values(u_combo, sorted(self.unit_values)))
        u_combo.bind("<FocusOut>", lambda _e: None)
        ttk.Label(r2, text=f"{self.tr('payment')}: ").pack(side="left")
        ttk.Entry(r2, textvariable=payment_var, width=14).pack(side="left", padx=(4, 0))

        r3 = ttk.Frame(w); r3.pack(fill="x", padx=10, pady=(0, 10))
        ttk.Label(r3, text=f"{self.tr('merchant')}: ").pack(side="left")
        ttk.Entry(r3, textvariable=merchant_var, width=22).pack(side="left", padx=(4, 10))
        ttk.Label(r3, text=f"{self.tr('category')}: ").pack(side="left")
        ttk.Entry(r3, textvariable=category_var, width=18).pack(side="left", padx=(4, 10))
        ttk.Checkbutton(r3, text=self.tr("excluded"), variable=excluded_var).pack(side="left")

        btn = ttk.Frame(w); btn.pack(side="bottom", fill="x", padx=10, pady=10)

        def apply_edit() -> None:
            try:
                iso_date = datetime.strptime(date_var.get().strip(), self.ui_date_fmt()).strftime("%Y-%m-%d")
            except ValueError:
                messagebox.showerror(self.tr("app_title"), self.tr("invalid_date"))
                return
            if not item_var.get().strip():
                messagebox.showerror(self.tr("app_title"), self.tr("invalid_item"))
                return
            try:
                raw = float(amount_var.get().strip())
            except ValueError:
                messagebox.showerror(self.tr("app_title"), self.tr("invalid_amount"))
                return
            raw = abs(raw)

            unit = unit_var.get().strip().upper()
            if unit not in self.unit_values:
                messagebox.showerror(self.tr("app_title"), self.tr("invalid_amount"))
                return

            selected_type = self._type_from_ui(type_var.get())
            amount = (-raw) if selected_type == "income" else raw

            self.entries[idx] = Entry(
                date=iso_date,
                amount=amount,
                item=item_var.get().strip(),
                unit=unit,
                payment=(payment_var.get().strip() or "-"),
                merchant=(merchant_var.get().strip() or "-"),
                category=(category_var.get().strip() or "-"),
                excluded=excluded_var.get(),
            )
            self._rebuild_options()
            self.write_current_delr()
            self._save_settings()
            w.destroy()
            self.refresh_table()

        ttk.Button(btn, text=self.tr("apply"), command=apply_edit).pack(side="right")
        ttk.Button(btn, text=self.tr("cancel"), command=w.destroy).pack(side="right", padx=(0, 6))
        self._finalize_filter_window(w, pad_w=30, pad_h=30)
    def delete_selected(self) -> None:
        ids = list(self.tree.selection())
        if not ids:
            messagebox.showerror(self.tr("app_title"), self.tr("no_select"))
            return
        for i in sorted([int(x) for x in ids], reverse=True):
            if 0 <= i < len(self.entries):
                del self.entries[i]
        self.write_current_delr(); self._save_settings(); self.refresh_table()

    def _fmt_number_ui(self, value: float) -> str:
        text = f"{value:.2f}"
        if self.code() == "de":
            return text.replace(".", ",")
        return text

    def _fmt(self, money: dict[str, float]) -> str:
        if not money:
            return self._fmt_number_ui(0.0)
        return " | ".join([f"{k}: {self._fmt_number_ui(v)}" for k, v in sorted(money.items())])

    def refresh_table(self) -> None:
        for iid in self.tree.get_children():
            self.tree.delete(iid)
        mode = self.view_mode_var.get()
        rows_all = list(enumerate(self.entries))

        if mode == "all":
            self.page_keys = []
            self.page_index = 0
            filtered_rows = [(i, e) for i, e in rows_all if self._passes_filters(e)]
            rows = self._sorted(filtered_rows)
        else:
            all_keys = sorted({self._make_page_key(e.date) for _, e in rows_all})
            self.page_keys = all_keys
            if not self.page_keys:
                self.page_index = 0
                rows = []
            else:
                if self.page_index >= len(self.page_keys):
                    self.page_index = len(self.page_keys) - 1
                key = self.page_keys[self.page_index]
                page_rows = [(i, e) for i, e in rows_all if self._make_page_key(e.date) == key]
                filtered_rows = [(i, e) for i, e in page_rows if self._passes_filters(e)]
                rows = self._sorted(filtered_rows)

        total: dict[str, float] = {}
        inc: dict[str, float] = {}
        exp: dict[str, float] = {}
        for idx, e in rows:
            u = (e.unit or "EUR").upper()
            a = abs(e.amount)
            if not e.excluded:
                total[u] = total.get(u, 0.0) + e.amount
                if e.amount < 0:
                    inc[u] = inc.get(u, 0.0) + a
                else:
                    exp[u] = exp.get(u, 0.0) + a
            display_type = self.tr("income") if e.amount < 0 else self.tr("expense")
            excluded = self.tr("excluded") if e.excluded else ""
            self.tree.insert("", "end", iid=str(idx), values=(self.fmt_ui_date(e.date), display_type, e.item, self._fmt_number_ui(a), u, e.payment, e.merchant, e.category, excluded))

        self.total_label.config(text=f"{self.tr('total')}: {self._fmt(total)}")
        self.income_label.config(text=f"{self.tr('income_total')}: {self._fmt(inc)}")
        self.expense_label.config(text=f"{self.tr('expense_total')}: {self._fmt(exp)}")
        self.update_header_state()
        self._update_paging_controls()
        self._update_action_buttons()


    def _create_date_picker(self, parent: ttk.Frame, vars3: tuple[tk.StringVar, tk.StringVar, tk.StringVar]) -> None:
        yv, mv, dv = vars3
        y_combo = ttk.Combobox(parent, textvariable=yv, values=[str(y) for y in range(2000, 2201)], width=6, state="readonly")
        m_combo = ttk.Combobox(parent, textvariable=mv, values=[f"{m:02d}" for m in range(1, 13)], width=4, state="readonly")
        d_combo = ttk.Combobox(parent, textvariable=dv, values=[f"{d:02d}" for d in range(1, 32)], width=4, state="readonly")

        def refresh_days(*_args) -> None:
            try:
                y = int(yv.get())
                m = int(mv.get())
                maxd = calendar.monthrange(y, m)[1]
            except Exception:
                maxd = 31
            vals = [f"{d:02d}" for d in range(1, maxd + 1)]
            d_combo["values"] = vals
            if dv.get() not in vals:
                dv.set(vals[-1])

        y_combo.bind("<<ComboboxSelected>>", refresh_days)
        m_combo.bind("<<ComboboxSelected>>", refresh_days)

        if self.code() == "en":
            order = [(m_combo, "M"), (d_combo, "D"), (y_combo, "Y")]
        elif self.code() == "de":
            order = [(d_combo, "T"), (m_combo, "M"), (y_combo, "J")]
        else:
            order = [(y_combo, "年"), (m_combo, "月"), (d_combo, "日")]

        for combo, label in order:
            combo.pack(side="left", padx=(4, 2))
            ttk.Label(parent, text=label).pack(side="left", padx=(0, 4))

        refresh_days()

    def _update_days(self) -> None:
        try:
            y = int(self.y_var.get())
            m = int(self.m_var.get())
            maxd = calendar.monthrange(y, m)[1]
        except Exception:
            maxd = 31
        vals = [f"{d:02d}" for d in range(1, maxd + 1)]
        self.d_combo["values"] = vals
        if self.d_var.get() not in vals:
            self.d_var.set(vals[-1])

    def _rebuild_options(self) -> None:
        self.payment_values = {e.payment for e in self.entries if e.payment and e.payment != "-"}
        self.category_values = {e.category for e in self.entries if e.category and e.category != "-"}
        self.merchant_values = {e.merchant for e in self.entries if e.merchant and e.merchant != "-"}
        self.unit_values = set(ISO_CURRENCIES) | {e.unit.upper() for e in self.entries if e.unit}

        payment_list = [self.cash()] + sorted(self.payment_values)
        self.payment_combo["values"] = payment_list
        self.category_combo["values"] = sorted(self.category_values)
        self.merchant_combo["values"] = sorted(self.merchant_values)
        self.unit_combo["values"] = sorted(self.unit_values)

    def read_entries_from_path(self, path: Path) -> list[Entry]:
        ext = path.suffix.lower()
        if ext in {".delr", ".csv"}:
            rows = self._read_rows_csv_like(path, delimiter=",")
        elif ext == ".tsv":
            rows = self._read_rows_csv_like(path, delimiter="	")
        elif ext == ".json":
            rows = self._read_rows_json(path)
        elif ext == ".xml":
            rows = self._read_rows_xml(path)
        elif ext in {".yaml", ".yml"}:
            rows = self._read_rows_yaml(path)
        elif ext == ".xlsx":
            rows = self._read_rows_xlsx(path)
        else:
            rows = []
        return self._rows_to_entries(rows)

    def write_current_delr(self) -> None:
        self._ensure_current()
        if self.current_file is None:
            return
        self.current_file.parent.mkdir(parents=True, exist_ok=True)
        with self.current_file.open("w", encoding="utf-8", newline="") as f:
            w = csv.DictWriter(f, fieldnames=CSV_HEADERS)
            w.writeheader()
            for e in self.entries:
                w.writerow({
                    "date": e.date,
                    "amount": f"{e.amount:.2f}",
                    "item": e.item,
                    "unit": e.unit,
                    "payment": e.payment,
                    "merchant": e.merchant,
                    "category": e.category,
                    "excluded": "1" if e.excluded else "0",
                })
        self.file_var.set(str(self.current_file))

    def new_ledger(self) -> None:
        folder = Path(self.folder_var.get().strip() or str(self.user_dir))
        folder.mkdir(parents=True, exist_ok=True)
        chosen = filedialog.asksaveasfilename(
            title=self.tr("new"),
            initialdir=str(folder),
            defaultextension=".delr",
            filetypes=[("DELR", "*.delr")],
        )
        if chosen:
            self.current_file = Path(chosen)
        else:
            self.current_file = folder / datetime.now().strftime("ledger-%Y%m%d-%H%M%S.delr")
        self.entries = []
        self._rebuild_options()
        self.write_current_delr()
        self._save_settings()
        self.refresh_table()

    def open_ledger(self) -> None:
        folder = Path(self.folder_var.get().strip() or str(self.user_dir))
        path = filedialog.askopenfilename(
            title=self.tr("open"),
            initialdir=str(folder),
            filetypes=[("Ledger", "*.delr *.csv *.tsv *.json *.xml *.yaml *.yml *.xlsx")],
        )
        if not path:
            return
        self.current_file = Path(path)
        try:
            self.entries = self.read_entries_from_path(self.current_file)
        except RuntimeError as e:
            messagebox.showerror(self.tr("app_title"), str(e))
            return
        self.file_var.set(str(self.current_file))
        self._rebuild_options()
        self._save_settings()
        self.refresh_table()

    def export_ledger(self) -> None:
        if not self.entries:
            return
        folder = Path(self.folder_var.get().strip() or str(self.user_dir))
        path = filedialog.asksaveasfilename(
            title=self.tr("export"),
            initialdir=str(folder),
            defaultextension=".delr",
            filetypes=[
                ("DELR", "*.delr"),
                ("CSV", "*.csv"),
                ("TSV", "*.tsv"),
                ("XLSX", "*.xlsx"),
                ("JSON", "*.json"),
                ("XML", "*.xml"),
                ("YAML", "*.yaml"),
            ],
        )
        if not path:
            return
        out = Path(path)
        ext = out.suffix.lower()
        rows = [{
            "date": e.date,
            "amount": f"{e.amount:.2f}",
            "item": e.item,
            "unit": e.unit,
            "payment": e.payment,
            "merchant": e.merchant,
            "category": e.category,
            "excluded": "1" if e.excluded else "0",
        } for e in self.entries]

        if ext in {".delr", ".csv"}:
            with out.open("w", encoding="utf-8", newline="") as f:
                w = csv.DictWriter(f, fieldnames=CSV_HEADERS)
                w.writeheader(); w.writerows(rows)
        elif ext == ".tsv":
            with out.open("w", encoding="utf-8", newline="") as f:
                w = csv.DictWriter(f, fieldnames=CSV_HEADERS, delimiter="	")
                w.writeheader(); w.writerows(rows)
        elif ext == ".json":
            out.write_text(json.dumps(rows, ensure_ascii=False, indent=2), encoding="utf-8")
        elif ext == ".xml":
            root = Element("ledger")
            for r in rows:
                node = SubElement(root, "entry")
                for k in CSV_HEADERS:
                    SubElement(node, k).text = r[k]
            ElementTree(root).write(out, encoding="utf-8", xml_declaration=True)
        elif ext in {".yaml", ".yml"}:
            if yaml is None:
                messagebox.showerror(self.tr("app_title"), "PyYAML is not installed")
                return
            out.write_text(yaml.safe_dump(rows, allow_unicode=True, sort_keys=False), encoding="utf-8")
        elif ext == ".xlsx":
            if Workbook is None:
                messagebox.showerror(self.tr("app_title"), "openpyxl is not installed")
                return
            wb = Workbook(); ws = wb.active
            ws.append(CSV_HEADERS)
            for r in rows:
                ws.append([r[k] for k in CSV_HEADERS])
            wb.save(out)

    def import_merge(self) -> None:
        folder = Path(self.folder_var.get().strip() or str(self.user_dir))
        path = filedialog.askopenfilename(
            title=self.tr("import"),
            initialdir=str(folder),
            filetypes=[("Ledger", "*.delr *.csv *.tsv *.json *.xml *.yaml *.yml *.xlsx")],
        )
        if not path:
            return
        try:
            more = self.read_entries_from_path(Path(path))
        except RuntimeError as e:
            messagebox.showerror(self.tr("app_title"), str(e))
            return
        if not more:
            return
        self.entries.extend(more)
        self._rebuild_options()
        self.write_current_delr()
        self._save_settings()
        self.refresh_table()

    def clear_form(self) -> None:
        self.item_var.set("")
        self.amount_var.set("")
        self.category_var.set("")
        self.merchant_var.set("")
        self.payment_var.set("")
        self.excluded_var.set(False)

    def add_entry(self) -> None:
        item = self.item_var.get().strip()
        if not item:
            messagebox.showerror(self.tr("app_title"), self.tr("invalid_item"))
            return
        try:
            y = int(self.y_var.get()); m = int(self.m_var.get()); d = int(self.d_var.get())
            iso_date = datetime(y, m, d).strftime("%Y-%m-%d")
        except Exception:
            messagebox.showerror(self.tr("app_title"), self.tr("invalid_date"))
            return
        try:
            raw = float(self.amount_var.get().strip())
        except ValueError:
            messagebox.showerror(self.tr("app_title"), self.tr("invalid_amount"))
            return
        raw = abs(raw)

        unit = self.unit_var.get().strip().upper()
        if unit not in self.unit_values:
            messagebox.showerror(self.tr("app_title"), self.tr("invalid_amount"))
            return

        selected_type = self._type_from_ui(self.type_var.get())
        amount = (-raw) if selected_type == "income" else raw

        payment = self.payment_var.get().strip() or "-"
        if payment in set(CASH_LABELS.values()):
            payment = self.cash()

        self.entries.append(Entry(
            date=iso_date,
            amount=amount,
            item=item,
            unit=unit,
            payment=payment,
            merchant=(self.merchant_var.get().strip() or "-"),
            category=(self.category_var.get().strip() or "-"),
            excluded=self.excluded_var.get(),
        ))
        self._rebuild_options()
        self.write_current_delr()
        self._save_settings()
        self.refresh_table()


    def _normalize_type_token(self, raw: str) -> str:
        return (raw or "").strip().casefold()

    def parse_entry_type(self, raw: str) -> str | None:
        token = self._normalize_type_token(raw)
        if token in {x.casefold() for x in ALL_TYPE_LABELS["income"]}:
            return "income"
        if token in {x.casefold() for x in ALL_TYPE_LABELS["expense"]}:
            return "expense"
        return None

    def parse_import_date(self, raw: str) -> str:
        t = (raw or "").strip()
        m = re.fullmatch(r"(\d{4})[-/.](\d{1,2})[-/.](\d{1,2})", t)
        if not m:
            raise ValueError("invalid date")
        y, mo, d = int(m.group(1)), int(m.group(2)), int(m.group(3))
        return datetime(y, mo, d).strftime("%Y-%m-%d")

    def parse_amount_value(self, raw: str) -> float:
        t = (raw or "").strip().replace(" ", "")
        if not t:
            raise ValueError("empty amount")
        sign = 1
        if t[0] in "+-":
            sign = -1 if t[0] == "-" else 1
            t = t[1:]
        if not t or not re.fullmatch(r"[0-9.,]+", t):
            raise ValueError("invalid amount")

        if "." in t and "," in t:
            last_dot = t.rfind(".")
            last_comma = t.rfind(",")
            dec_sep = "." if last_dot > last_comma else ","
            oth_sep = "," if dec_sep == "." else "."
            left, right = t.rsplit(dec_sep, 1)
            if right and len(right) in (1, 2):
                num = left.replace(oth_sep, "").replace(dec_sep, "") + "." + right
                return sign * float(num)
            num = t.replace(".", "").replace(",", "")
            return sign * float(num)

        sep = "." if "." in t else ("," if "," in t else "")
        if not sep:
            return sign * float(t)

        parts = t.split(sep)
        if len(parts) == 1:
            return sign * float(parts[0])
        last = parts[-1]
        if len(last) in (1, 2):
            num = "".join(parts[:-1]) + "." + last
            return sign * float(num)
        if len(last) == 3:
            num = "".join(parts)
            return sign * float(num)
        raise ValueError("invalid amount precision")

    def _try_match_unit(self, raw: str) -> str | None:
        token = (raw or "").strip().upper()
        if token and token in self.unit_values:
            return token
        return None

    def _try_match_payment(self, raw: str) -> bool:
        token = " ".join((raw or "").strip().casefold().split())
        return token in {x.casefold() for x in PAYMENT_HINTS}

    def _split_import_line(self, line: str) -> list[str]:
        text = line.strip()
        for sep in ("	", ";", "|"):
            parts = [p.strip() for p in text.split(sep)]
            non_empty = [p for p in parts if p]
            if len(non_empty) >= 2:
                return non_empty
        return [text] if text else []

    def _ask_middle_field_role(self, raw_text: str) -> str | None:
        w = tk.Toplevel(self.root)
        w.title(self.tr("choose_field_role"))
        w.resizable(False, False)
        w.transient(self.root)
        w.grab_set()
        ttk.Label(w, text=f"{raw_text}").pack(fill="x", padx=12, pady=(12, 8))
        result = {"role": None}
        row = ttk.Frame(w)
        row.pack(fill="x", padx=12, pady=(0, 12))

        def pick(role: str) -> None:
            result["role"] = role
            w.destroy()

        ttk.Button(row, text=self.tr("as_payment"), command=lambda: pick("payment")).pack(side="left")
        ttk.Button(row, text=self.tr("as_merchant"), command=lambda: pick("merchant")).pack(side="left", padx=(6, 0))
        self._finalize_filter_window(w)
        self.root.wait_window(w)
        return result["role"]

    def _show_import_errors(self, errors: list[str]) -> None:
        if not errors:
            return
        w = tk.Toplevel(self.root)
        w.title(self.tr("import_errors"))
        w.geometry("760x420")
        box = tk.Text(w, wrap="word")
        box.pack(fill="both", expand=True, padx=10, pady=10)
        box.insert("1.0", "\n".join(errors))
        box.config(state="disabled")

    def open_paste_import_window(self) -> None:
        w = tk.Toplevel(self.root)
        w.title(self.tr("paste_import"))
        w.geometry("860x520")
        w.transient(self.root)
        w.grab_set()

        help_text = f"{self.tr('import_help')}\n{self.tr('import_date_rule')}"
        ttk.Label(w, text=help_text, justify="left", wraplength=820).pack(fill="x", padx=12, pady=(12, 8))

        text = tk.Text(w, wrap="none")
        text.pack(fill="both", expand=True, padx=12, pady=(0, 10))

        try:
            clip = self.root.clipboard_get()
            if clip:
                text.insert("1.0", clip)
        except Exception:
            pass

        btn = ttk.Frame(w)
        btn.pack(fill="x", padx=12, pady=(0, 12))

        def do_import() -> None:
            self._import_from_pasted_text(text.get("1.0", "end"), w)

        ttk.Button(btn, text=self.tr("parse_and_import"), command=do_import).pack(side="right")
        ttk.Button(btn, text=self.tr("cancel"), command=w.destroy).pack(side="right", padx=(0, 6))

    def _import_from_pasted_text(self, raw: str, window: tk.Toplevel | None = None) -> None:
        lines = raw.splitlines()
        ok_entries: list[Entry] = []
        errors: list[str] = []

        last_unit = self.unit_var.get().strip().upper() or "EUR"
        last_payment = self.payment_var.get().strip()

        for ln, line in enumerate(lines, start=1):
            if not line.strip():
                continue
            try:
                fields = self._split_import_line(line)
                if len(fields) < 2:
                    raise ValueError("unsupported delimiter")

                date_val: str | None = None
                type_val: str | None = None
                amount_raw: float | None = None
                unit_val: str | None = None

                remaining: list[str] = []
                for f in fields:
                    captured = False

                    if date_val is None:
                        d: str | None = None
                        try:
                            d = self.parse_import_date(f)
                        except Exception:
                            d = None
                        if d is not None:
                            date_val = d
                            captured = True

                    if not captured and type_val is None:
                        t = self.parse_entry_type(f)
                        if t is not None:
                            type_val = t
                            captured = True

                    if not captured and amount_raw is None:
                        a: float | None = None
                        try:
                            a = self.parse_amount_value(f)
                        except Exception:
                            a = None
                        if a is not None:
                            amount_raw = a
                            captured = True

                    if not captured and unit_val is None:
                        u = self._try_match_unit(f)
                        if u is not None:
                            unit_val = u
                            captured = True

                    if not captured:
                        remaining.append(f)

                payment_val: str | None = None
                next_remaining: list[str] = []
                for f in remaining:
                    if payment_val is None and self._try_match_payment(f):
                        payment_val = f.strip()
                    else:
                        next_remaining.append(f)
                remaining = next_remaining

                item = ""
                merchant = "-"
                category = ""

                if len(remaining) == 1:
                    item = remaining[0].strip()
                elif len(remaining) == 2:
                    item, category = remaining[0].strip(), remaining[1].strip()
                elif len(remaining) == 3:
                    item = remaining[0].strip()
                    middle = remaining[1].strip()
                    category = remaining[2].strip()
                    if payment_val is None:
                        role = self._ask_middle_field_role(middle)
                        if role is None:
                            raise ValueError("role not selected")
                        if role == "payment":
                            payment_val = middle
                        else:
                            merchant = middle
                    else:
                        merchant = middle
                elif len(remaining) == 4:
                    item = remaining[0].strip()
                    payment_val = payment_val or remaining[1].strip()
                    merchant = remaining[2].strip()
                    category = remaining[3].strip()
                elif len(remaining) > 4:
                    raise ValueError("too many text fields")

                if date_val is None:
                    raise ValueError("missing date")
                if amount_raw is None:
                    raise ValueError("missing amount")
                if not item:
                    raise ValueError("missing item")
                if not category:
                    raise ValueError("missing category")

                unit_final = (unit_val or last_unit or "EUR").upper()
                if unit_final not in self.unit_values:
                    unit_final = "EUR"

                payment_final = (payment_val or last_payment or "-").strip() or "-"

                if type_val == "income":
                    amount_final = -amount_raw
                elif type_val == "expense":
                    amount_final = amount_raw
                else:
                    amount_final = amount_raw

                ok_entries.append(Entry(
                    date=date_val,
                    amount=amount_final,
                    item=item,
                    unit=unit_final,
                    payment=payment_final,
                    merchant=(merchant.strip() or "-"),
                    category=category,
                ))

                last_unit = unit_final
                if payment_final and payment_final != "-":
                    last_payment = payment_final
            except Exception as e:
                errors.append(f"Line {ln}: {line}\n  -> {e}")

        if ok_entries:
            self.entries.extend(ok_entries)
            self.unit_var.set(last_unit)
            self.payment_var.set(last_payment)
            self._rebuild_options()
            self.write_current_delr()
            self._save_settings()
            self.refresh_table()

        if window is not None:
            window.destroy()

        messagebox.showinfo(self.tr("paste_import"), self.tr("import_summary").format(ok=len(ok_entries), bad=len(errors)))
        if errors:
            self._show_import_errors(errors)



    def _get_view_rows(self, current_page_only: bool) -> list[tuple[int, Entry]]:
        rows_all = list(enumerate(self.entries))
        mode = self.view_mode_var.get()

        if mode == "all":
            filtered_rows = [(i, e) for i, e in rows_all if self._passes_filters(e)]
            return self._sorted(filtered_rows)

        all_keys = sorted({self._make_page_key(e.date) for _, e in rows_all})
        if not all_keys:
            return []

        def page_rows_for(key: str) -> list[tuple[int, Entry]]:
            page_rows = [(i, e) for i, e in rows_all if self._make_page_key(e.date) == key]
            filtered_rows = [(i, e) for i, e in page_rows if self._passes_filters(e)]
            return self._sorted(filtered_rows)

        if current_page_only:
            idx = min(max(self.page_index, 0), len(all_keys) - 1)
            return page_rows_for(all_keys[idx])

        out: list[tuple[int, Entry]] = []
        for k in all_keys:
            out.extend(page_rows_for(k))
        return out

    def _format_document_amount(self, amount: float, unit: str) -> str:
        return f"{self._fmt_number_ui(amount)} {unit}"

    def _group_rows_by_date(self, rows: list[tuple[int, Entry]]) -> OrderedDict[str, list[Entry]]:
        grouped: OrderedDict[str, list[Entry]] = OrderedDict()
        for _i, e in rows:
            grouped.setdefault(e.date, []).append(e)
        return grouped

    def _get_doc_headers(self) -> tuple[str, str, str, str, str]:
        return (
            self.tr("doc_item"),
            self.tr("doc_price"),
            self.tr("doc_payment"),
            self.tr("doc_merchant"),
            self.tr("doc_category"),
        )

    def _choose_document_format(self) -> str | None:
        w = tk.Toplevel(self.root)
        w.title(self.tr("choose_doc_format"))
        w.resizable(False, False)
        w.transient(self.root)
        w.grab_set()

        ttk.Label(w, text=self.tr("choose_doc_format")).pack(fill="x", padx=14, pady=(14, 10))

        result = {"fmt": None}

        def choose(fmt: str) -> None:
            result["fmt"] = fmt
            w.destroy()

        row = ttk.Frame(w)
        row.pack(fill="x", padx=14, pady=(0, 14))
        ttk.Button(row, text=self.tr("doc_format_md"), command=lambda: choose("md")).pack(fill="x", pady=(0, 6))
        ttk.Button(row, text=self.tr("doc_format_docx"), command=lambda: choose("docx")).pack(fill="x", pady=(0, 6))
        ttk.Button(row, text=self.tr("doc_format_pdf"), command=lambda: choose("pdf")).pack(fill="x")

        self._finalize_filter_window(w, pad_w=40, pad_h=20)
        self.root.wait_window(w)
        return result["fmt"]

    def export_document(self) -> None:
        fmt = self._choose_document_format()
        if fmt is None:
            return

        rows = self._get_view_rows(self.doc_current_page_only_var.get())
        if not rows:
            messagebox.showinfo(self.tr("app_title"), self.tr("no_select"))
            return

        folder = Path(self.folder_var.get().strip() or str(self.user_dir))
        ext = f".{fmt}"
        path = filedialog.asksaveasfilename(
            title=self.tr("export_doc"),
            initialdir=str(folder),
            defaultextension=ext,
            filetypes=[(fmt.upper(), f"*{ext}")],
        )
        if not path:
            return
        out = Path(path)

        grouped = self._group_rows_by_date(rows)
        if fmt == "md":
            self.export_document_markdown(out, grouped)
        elif fmt == "docx":
            self.export_document_docx(out, grouped)
        elif fmt == "pdf":
            self.export_document_pdf(out, grouped)

    def export_document_markdown(self, out: Path, grouped: OrderedDict[str, list[Entry]]) -> None:
        h_item, h_price, h_payment, h_merchant, h_category = self._get_doc_headers()
        lines: list[str] = [f"# {self.tr('doc_title')}", ""]
        grand = 0.0

        for d, items in grouped.items():
            lines.append(f"## {self.fmt_ui_date(d)}")
            lines.append("")
            lines.append(f"| {h_item} | {h_price} | {h_payment} | {h_merchant} | {h_category} |")
            lines.append("| ---- | ----: | ------- | -------- | -------- |")
            sub = 0.0
            for e in items:
                if not e.excluded:
                    sub += e.amount
                    grand += e.amount
                lines.append(f"| {e.item} | {self._format_document_amount(e.amount, e.unit.upper())} | {e.payment} | {e.merchant} | {e.category} |")
            lines.append("")
            unit = items[0].unit.upper() if items else "EUR"
            lines.append(f"**{self.tr('doc_subtotal')}: {self._format_document_amount(sub, unit)}**")
            lines.append("")

        total_unit = next(iter(grouped.values()))[0].unit.upper() if grouped else "EUR"
        lines.append(f"**{self.tr('doc_total')}: {self._format_document_amount(grand, total_unit)}**")
        out.write_text("\n".join(lines), encoding="utf-8")

    def export_document_docx(self, out: Path, grouped: OrderedDict[str, list[Entry]]) -> None:
        if Document is None:
            messagebox.showerror(self.tr("app_title"), "python-docx is not installed")
            return

        h_item, h_price, h_payment, h_merchant, h_category = self._get_doc_headers()
        doc = Document()
        doc.add_heading(self.tr("doc_title"), level=1)
        grand = 0.0

        for d, items in grouped.items():
            doc.add_heading(self.fmt_ui_date(d), level=2)
            table = doc.add_table(rows=1, cols=5)
            hdr = table.rows[0].cells
            hdr[0].text = h_item
            hdr[1].text = h_price
            hdr[2].text = h_payment
            hdr[3].text = h_merchant
            hdr[4].text = h_category
            sub = 0.0
            for e in items:
                if not e.excluded:
                    sub += e.amount
                    grand += e.amount
                row = table.add_row().cells
                row[0].text = e.item
                row[1].text = self._format_document_amount(e.amount, e.unit.upper())
                row[2].text = e.payment
                row[3].text = e.merchant
                row[4].text = e.category
            unit = items[0].unit.upper() if items else "EUR"
            doc.add_paragraph(f"{self.tr('doc_subtotal')}: {self._format_document_amount(sub, unit)}")

        total_unit = next(iter(grouped.values()))[0].unit.upper() if grouped else "EUR"
        doc.add_paragraph(f"{self.tr('doc_total')}: {self._format_document_amount(grand, total_unit)}")
        doc.save(str(out))

    def _get_pdf_font_name(self) -> str | None:
        if pdfmetrics is None or TTFont is None:
            return None

        candidates = [
            ("DELRPdfCJK", Path(r"C:\Windows\Fonts\simhei.ttf")),
            ("DELRPdfCJK", Path(r"C:\Windows\Fonts\simsunb.ttf")),
            ("DELRPdfCJK", Path(r"C:\Windows\Fonts\NotoSansTC-VF.ttf")),
            ("DELRPdfCJK", Path(r"C:\Windows\Fonts\NotoSerifTC-VF.ttf")),
        ]

        registered = set(pdfmetrics.getRegisteredFontNames())
        for font_name, font_path in candidates:
            if font_name in registered:
                return font_name
            if font_path.exists():
                try:
                    pdfmetrics.registerFont(TTFont(font_name, str(font_path)))
                    return font_name
                except Exception:
                    continue
        return None

    def _contains_cjk(self, text: str) -> bool:
        return any(
            (0x4E00 <= ord(ch) <= 0x9FFF)
            or (0x3400 <= ord(ch) <= 0x4DBF)
            or (0x3000 <= ord(ch) <= 0x303F)
            or (0xFE30 <= ord(ch) <= 0xFE4F)
            or (0xFF01 <= ord(ch) <= 0xFF5E)
            for ch in text
        )

    def _pdf_paragraph(self, text: str, style, cjk_font: str | None):
        if not cjk_font or not text:
            return Paragraph(escape(text), style)

        parts: list[str] = []
        current: list[str] = []
        current_is_cjk: bool | None = None

        for ch in text:
            is_cjk = self._contains_cjk(ch)
            if current_is_cjk is None or is_cjk == current_is_cjk:
                current.append(ch)
                current_is_cjk = is_cjk
                continue

            chunk = escape("".join(current))
            if current_is_cjk:
                parts.append(f'<font name="{cjk_font}">{chunk}</font>')
            else:
                parts.append(chunk)
            current = [ch]
            current_is_cjk = is_cjk

        if current:
            chunk = escape("".join(current))
            if current_is_cjk:
                parts.append(f'<font name="{cjk_font}">{chunk}</font>')
            else:
                parts.append(chunk)

        return Paragraph("".join(parts), style)

    def export_document_pdf(self, out: Path, grouped: OrderedDict[str, list[Entry]]) -> None:
        if SimpleDocTemplate is None or getSampleStyleSheet is None or Table is None or TableStyle is None or colors is None or A4 is None:
            messagebox.showerror(self.tr("app_title"), "reportlab is not installed")
            return

        styles = getSampleStyleSheet()
        font_name = self._get_pdf_font_name()
        story = [self._pdf_paragraph(self.tr("doc_title"), styles["Title"], font_name), Spacer(1, 8)]
        h_item, h_price, h_payment, h_merchant, h_category = self._get_doc_headers()
        grand = 0.0

        for d, items in grouped.items():
            story.append(self._pdf_paragraph(self.fmt_ui_date(d), styles["Heading2"], font_name))
            data = [[
                self._pdf_paragraph(h_item, styles["Normal"], font_name),
                self._pdf_paragraph(h_price, styles["Normal"], font_name),
                self._pdf_paragraph(h_payment, styles["Normal"], font_name),
                self._pdf_paragraph(h_merchant, styles["Normal"], font_name),
                self._pdf_paragraph(h_category, styles["Normal"], font_name),
            ]]
            sub = 0.0
            for e in items:
                if not e.excluded:
                    sub += e.amount
                    grand += e.amount
                data.append([
                    self._pdf_paragraph(e.item, styles["Normal"], font_name),
                    self._pdf_paragraph(self._format_document_amount(e.amount, e.unit.upper()), styles["Normal"], font_name),
                    self._pdf_paragraph(e.payment, styles["Normal"], font_name),
                    self._pdf_paragraph(e.merchant, styles["Normal"], font_name),
                    self._pdf_paragraph(e.category, styles["Normal"], font_name),
                ])
            t = Table(data, repeatRows=1)
            t.setStyle(TableStyle([
                ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
                ("BACKGROUND", (0, 0), (-1, 0), colors.lightgrey),
                ("ALIGN", (1, 1), (1, -1), "RIGHT"),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("FONTSIZE", (0, 0), (-1, -1), 10),
            ]))
            story.append(t)
            unit = items[0].unit.upper() if items else "EUR"
            story.append(Spacer(1, 4))
            story.append(self._pdf_paragraph(f"{self.tr('doc_subtotal')}: {self._format_document_amount(sub, unit)}", styles["Normal"], font_name))
            story.append(Spacer(1, 8))

        total_unit = next(iter(grouped.values()))[0].unit.upper() if grouped else "EUR"
        story.append(self._pdf_paragraph(f"{self.tr('doc_total')}: {self._format_document_amount(grand, total_unit)}", styles["Heading3"], font_name))
        pdf = SimpleDocTemplate(str(out), pagesize=A4)
        pdf.build(story)

def runtime_app_dir() -> Path:
    return Path(sys.executable).resolve().parent if getattr(sys, "frozen", False) else Path(__file__).resolve().parent


def main() -> None:
    root = tk.Tk()
    DelrLedgerApp(root, runtime_app_dir())
    root.mainloop()


if __name__ == "__main__":
    main()

























































